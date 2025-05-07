#!/usr/bin/env python
# -*- coding: utf-8 -*-

import os
from werkzeug.utils import secure_filename
from flask import current_app
import uuid

# Поддерживаемые расширения файлов
ALLOWED_EXTENSIONS = {'pdf', 'doc', 'docx', 'jpg', 'jpeg', 'png'}

def allowed_file(filename):
    """Проверка допустимости расширения файла"""
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def save_resume(file, tracking_code):
    """Сохранение файла резюме"""
    if not file or not allowed_file(file.filename):
        return None
    
    # Получаем расширение файла
    filename = secure_filename(file.filename)
    
    # Проверяем наличие расширения
    parts = filename.rsplit('.', 1) if '.' in filename else [filename, '']
    if len(parts) < 2 or not parts[1]:
        # Если расширение отсутствует или пустое, используем расширение по умолчанию
        extension = 'pdf'  # Расширение по умолчанию
    else:
        extension = parts[1].lower()
    
    # Создаем имя файла на основе кода отслеживания
    new_filename = f"resume_{tracking_code}.{extension}"
    
    # Определяем путь для сохранения
    upload_folder = current_app.config['UPLOAD_FOLDER']
    
    # Создаем директорию, если она не существует
    os.makedirs(upload_folder, exist_ok=True)
    
    # Путь к файлу
    file_path = os.path.join(upload_folder, new_filename)
    
    # Сохраняем файл
    file.save(file_path)
    
    return file_path

def extract_text_from_resume(file_path):
    """Извлечение текста из файла резюме"""
    current_app.logger.info(f"Начало извлечения текста из файла: {file_path}")
    
    if not file_path or not os.path.exists(file_path):
        current_app.logger.error(f"Файл не существует или путь не указан: {file_path}")
        return None
    
    # Проверяем наличие расширения
    parts = file_path.rsplit('.', 1) if '.' in file_path else [file_path, '']
    if len(parts) < 2 or not parts[1]:
        # Если расширение отсутствует или пустое, используем расширение по умолчанию
        extension = 'pdf'  # Расширение по умолчанию
        current_app.logger.info(f"Расширение не найдено, используем по умолчанию: {extension}")
    else:
        extension = parts[1].lower()
        current_app.logger.info(f"Определено расширение файла: {extension}")
    
    # Извлечение текста из разных форматов файлов
    if extension == 'pdf':
        current_app.logger.info("Извлечение текста из PDF")
        return extract_text_from_pdf(file_path)
    elif extension in ['doc', 'docx']:
        current_app.logger.info("Извлечение текста из DOC/DOCX")
        return extract_text_from_docx(file_path)
    elif extension in ['jpg', 'jpeg', 'png']:
        current_app.logger.info("Извлечение текста из изображения")
        return extract_text_from_image(file_path)
    
    current_app.logger.error(f"Неподдерживаемый формат файла: {extension}")
    return None

def extract_text_from_pdf(file_path):
    """Извлечение текста из PDF-файла"""
    try:
        import fitz  # PyMuPDF
        current_app.logger.info("PyMuPDF успешно импортирован")
        
        text = ""
        with fitz.open(file_path) as pdf:
            current_app.logger.info(f"PDF файл открыт, количество страниц: {len(pdf)}")
            for page_num, page in enumerate(pdf, 1):
                page_text = page.get_text()
                text += page_text
                current_app.logger.info(f"Извлечен текст со страницы {page_num}, длина: {len(page_text)} символов")
        
        if not text.strip():
            current_app.logger.warning("Извлеченный текст пуст")
            return None
            
        current_app.logger.info(f"Успешно извлечен текст из PDF, общая длина: {len(text)} символов")
        return text
    except ImportError:
        current_app.logger.error("PyMuPDF (fitz) не установлен. Установите: pip install pymupdf")
        return "Извлечение текста из PDF недоступно. Установите необходимые библиотеки."
    except Exception as e:
        current_app.logger.error(f"Ошибка при извлечении текста из PDF: {str(e)}", exc_info=True)
        return None

def extract_text_from_docx(file_path):
    """Извлечение текста из DOCX/DOC файла"""
    try:
        import docx
        current_app.logger.info("python-docx успешно импортирован")
        
        if file_path.endswith('.docx'):
            doc = docx.Document(file_path)
            current_app.logger.info(f"DOCX файл открыт, количество параграфов: {len(doc.paragraphs)}")
            
            text = "\n".join([para.text for para in doc.paragraphs])
            
            if not text.strip():
                current_app.logger.warning("Извлеченный текст пуст")
                return None
                
            current_app.logger.info(f"Успешно извлечен текст из DOCX, длина: {len(text)} символов")
            return text
        else:
            current_app.logger.warning("Извлечение текста из DOC-файлов в данный момент не поддерживается")
            return "Извлечение текста из DOC-файлов в данный момент не поддерживается"
    except ImportError:
        current_app.logger.error("python-docx не установлен. Установите: pip install python-docx")
        return "Извлечение текста из DOCX недоступно. Установите необходимые библиотеки."
    except Exception as e:
        current_app.logger.error(f"Ошибка при извлечении текста из DOCX: {str(e)}", exc_info=True)
        return None

def extract_text_from_image(file_path):
    """Извлечение текста из изображения (OCR) - временно отключено"""
    current_app.logger.info(f"Получен запрос на извлечение текста из изображения: {file_path}")
    current_app.logger.info("Функция извлечения текста из изображений временно отключена")
    
    # Возвращаем стандартный текст, указывая, что это изображение
    return "Это изображение. Извлечение текста из изображений временно отключено. Оригинальный файл сохранен и доступен для скачивания." 